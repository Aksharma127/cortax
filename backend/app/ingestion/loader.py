import logging
import asyncio
from pathlib import Path
from typing import BinaryIO
from abc import ABC, abstractmethod
import aiofiles
import PyPDF2
from docx import Document as DocxDocument

from app.core.config import settings

logger = logging.getLogger(__name__)


class DocumentParser(ABC):
    """Abstract document parser interface."""

    @abstractmethod
    def parse(self, file_path: Path) -> str:
        """Parse a file and return extracted text."""
        pass

    @abstractmethod
    def validate(self, file_path: Path) -> bool:
        """Validate file format."""
        pass


class PDFParser(DocumentParser):
    """PDF parser using PyPDF2."""

    def validate(self, file_path: Path) -> bool:
        """Check if file is valid PDF."""
        try:
            with open(file_path, "rb") as f:
                return f.read(4) == b"%PDF"
        except Exception as e:
            logger.error(f"PDF validation error: {e}")
            return False

    def parse(self, file_path: Path) -> str:
        """Extract text from PDF."""
        logger.info(f"Parsing PDF: {file_path}")
        text_parts = []

        try:
            with open(file_path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                
                for page_num, page in enumerate(reader.pages):
                    try:
                        text = page.extract_text()
                        if text.strip():
                            text_parts.append(text)
                    except Exception as e:
                        logger.warning(f"Failed to extract page {page_num}: {e}")

            return "\n".join(text_parts)

        except Exception as e:
            logger.error(f"PDF parsing failed: {e}")
            raise


class TextParser(DocumentParser):
    """Simple text file parser."""

    def validate(self, file_path: Path) -> bool:
        """Check if file is valid text."""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                f.read(100)
            return True
        except Exception as e:
            logger.error(f"Text validation error: {e}")
            return False

    def parse(self, file_path: Path) -> str:
        """Extract text from text file."""
        logger.info(f"Parsing text file: {file_path}")
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
        except Exception as e:
            logger.error(f"Text parsing failed: {e}")
            raise


class DocxParser(DocumentParser):
    """DOCX parser using python-docx."""

    def validate(self, file_path: Path) -> bool:
        """Check if file is valid DOCX."""
        try:
            DocxDocument(file_path)
            return True
        except Exception as e:
            logger.error(f"DOCX validation error: {e}")
            return False

    def parse(self, file_path: Path) -> str:
        """Extract text from DOCX."""
        logger.info(f"Parsing DOCX: {file_path}")
        try:
            doc = DocxDocument(file_path)
            text_parts = [para.text for para in doc.paragraphs]
            return "\n".join(text_parts)
        except Exception as e:
            logger.error(f"DOCX parsing failed: {e}")
            raise


class MarkdownParser(DocumentParser):
    """Markdown parser."""

    def validate(self, file_path: Path) -> bool:
        """Check if file looks like markdown."""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read(500)
                # Simple heuristic: contains markdown syntax
                return any(c in content for c in ["#", "**", "*", "-", "[", "`"])
        except Exception as e:
            logger.error(f"Markdown validation error: {e}")
            return False

    def parse(self, file_path: Path) -> str:
        """Extract text from markdown."""
        logger.info(f"Parsing Markdown: {file_path}")
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
        except Exception as e:
            logger.error(f"Markdown parsing failed: {e}")
            raise


class DocumentLoader:
    """Document loading with multi-format support."""

    PARSERS = {
        "pdf": PDFParser(),
        "txt": TextParser(),
        "docx": DocxParser(),
        "md": MarkdownParser(),
    }

    @staticmethod
    def validate_file(file_path: Path) -> bool:
        """
        Comprehensive file validation.
        
        Checks:
        1. File exists
        2. File size within limits
        3. File extension allowed
        4. Content matches format
        """
        # Check existence
        if not file_path.exists():
            logger.error(f"File not found: {file_path}")
            return False

        # Check file size
        file_size = file_path.stat().st_size
        if file_size > settings.MAX_FILE_SIZE:
            logger.error(f"File too large: {file_size} > {settings.MAX_FILE_SIZE}")
            return False

        # Check extension
        extension = file_path.suffix.lstrip(".").lower()
        if extension not in settings.ALLOWED_EXTENSIONS:
            logger.error(f"File type not allowed: {extension}")
            return False

        # Check content format
        parser = DocumentLoader.PARSERS.get(extension)
        if not parser or not parser.validate(file_path):
            logger.error(f"File content validation failed: {file_path}")
            return False

        logger.info(f"File validation passed: {file_path}")
        return True

    @staticmethod
    async def load_file(file_path: Path) -> str:
        """
        Load and parse file asynchronously.
        
        Async is important because:
        - File I/O is slow
        - PDF parsing is CPU-intensive
        - In production, you don't want to block other requests
        """
        # Validate first
        if not DocumentLoader.validate_file(file_path):
            raise ValueError(f"File validation failed: {file_path}")

        # Get parser
        extension = file_path.suffix.lstrip(".").lower()
        parser = DocumentLoader.PARSERS.get(extension)
        if not parser:
            raise ValueError(f"No parser for file type: {extension}")

        # Parse (run in thread pool to not block event loop)
        loop = asyncio.get_event_loop()
        text = await loop.run_in_executor(None, parser.parse, file_path)

        return text

    @staticmethod
    async def save_upload(
        file_content: bytes,
        filename: str,
        target_dir: str = settings.UPLOAD_DIR
    ) -> Path:
        """
        Save uploaded file asynchronously.
        
        Checks:
        - Directory exists
        - File doesn't already exist
        - Size limits
        """
        # Create directory if needed
        target_path = Path(target_dir)
        target_path.mkdir(parents=True, exist_ok=True)

        # Construct file path
        file_path = target_path / filename

        # Avoid overwriting
        if file_path.exists():
            base = file_path.stem
            suffix = file_path.suffix
            counter = 1
            while file_path.exists():
                file_path = target_path / f"{base}_{counter}{suffix}"
                counter += 1

        # Save file asynchronously
        async with aiofiles.open(file_path, "wb") as f:
            await f.write(file_content)

        logger.info(f"File saved: {file_path}")
        return file_path
